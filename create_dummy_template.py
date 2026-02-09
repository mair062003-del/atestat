from docx import Document
from docx.shared import Pt
import os

def create_dummy_template():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    doc.add_heading('ATTESTAT / АТТЕСТАТ', 0)
    
    doc.add_paragraph('This is a dummy template generated for demonstration.')
    doc.add_paragraph('Бұл демонстрацияға арналған үлгі.')
    
    p = doc.add_paragraph()
    p.add_run('Student / Оқушы: ').bold = True
    p.add_run('{{ student_name_kz }} / {{ student_name_ru }}')
    
    doc.add_heading('Grades / Бағалар', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subject (KZ)'
    hdr_cells[1].text = 'Subject (RU)'
    hdr_cells[2].text = 'Grade'
    
    # Add template loop row
    row = table.add_row().cells
    row[0].text = '{{ s.name_kz }}'
    row[1].text = '{{ s.name_ru }}'
    row[2].text = '{{ s.grade }}'
    
    # Jinja2 tags for loop
    # In docxtpl, loops are usually handled by surrounding the row with tags
    # But for simplicity with simple    # Add a simple text loop to demonstrate data injection without 'tr' tag complexity
    doc.add_heading('Subjects List:', level=2)
    p = doc.add_paragraph()
    # Just print the raw list to prove data is passing through.
    # The user will create a proper table in Word anyway.
    p.add_run("Subjects Data: {{ subjects }}")
    
    # Ensure directory exists
    output_dir = os.path.join(os.path.dirname(__file__), 'templates')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_path = os.path.join(output_dir, 'template.docx')
    doc.save(output_path)
    print(f"Dummy template created at '{output_path}'")

if __name__ == "__main__":
    create_dummy_template()
